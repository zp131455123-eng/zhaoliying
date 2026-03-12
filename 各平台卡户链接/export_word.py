#!/usr/bin/env python3
# -*- coding: utf-8 -*-
import json, os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DATA_FILE = os.path.join(os.path.dirname(__file__), "links.json")
OUT_FILE  = os.path.join(os.path.dirname(__file__), "各平台卡户链接.docx")

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), 'CCCCCC')
        tcBorders.append(border)
    tcPr.append(tcBorders)

with open(DATA_FILE, encoding='utf-8') as f:
    data = json.load(f)

doc = Document()

# 页边距
section = doc.sections[0]
section.left_margin   = Cm(2)
section.right_margin  = Cm(2)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)

# 标题
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('各平台开户客户 & 资金')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

doc.add_paragraph()

PLATFORM_COLOR = '1A3A5C'   # 深蓝
HEADER_COLOR   = '2E6DA4'   # 中蓝
ROW_ODD        = 'EEF4FB'   # 浅蓝
ROW_EVEN       = 'FFFFFF'   # 白

grand_total = 0

for d in data:
    clients = d.get('clients', [])
    if not clients:
        continue

    platform = d['platform']
    atype    = d['account_type']
    total    = sum(c['amount'] for c in clients)
    grand_total += total

    # 平台标题
    ph = doc.add_paragraph()
    run = ph.add_run(f'▶  {platform} / {atype}')
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    ph.paragraph_format.space_before = Pt(10)
    ph.paragraph_format.space_after  = Pt(4)

    # 表格：序号 | 姓名 | 资金(U)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    tbl.autofit = False
    tbl.columns[0].width = Cm(2)
    tbl.columns[1].width = Cm(7)
    tbl.columns[2].width = Cm(4)

    # 表头
    hdr = tbl.rows[0].cells
    for cell, txt in zip(hdr, ['序号', '姓名', '资金 (U)']):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        cell.paragraphs[0].runs[0].font.size = Pt(11)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, HEADER_COLOR)
        set_cell_border(cell)

    # 数据行
    for idx, c in enumerate(clients, 1):
        row = tbl.add_row().cells
        bg = ROW_ODD if idx % 2 == 1 else ROW_EVEN
        for cell, txt, align in zip(row,
            [str(idx), c['name'], f"{c['amount']:,}"],
            [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT]):
            cell.text = txt
            cell.paragraphs[0].runs[0].font.size = Pt(11)
            cell.paragraphs[0].alignment = align
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            set_cell_border(cell)

    # 合计行
    row = tbl.add_row().cells
    for cell, txt, align in zip(row,
        ['', '合　计', f"{total:,}"],
        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.RIGHT]):
        cell.text = txt
        r = cell.paragraphs[0].runs[0]
        r.font.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        cell.paragraphs[0].alignment = align
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, PLATFORM_COLOR)
        set_cell_border(cell)

    doc.add_paragraph()

# 总计
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f'💰 三平台总计：{grand_total:,} U')
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

doc.save(OUT_FILE)
print(f'已导出: {OUT_FILE}')
